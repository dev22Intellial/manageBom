import os
import openpyxl
import csv
from docx import Document
# from PyPDF2 import PdfReader # No longer used
import pdfplumber

class BaseBOMParser:
    """Base class for all BOM parsers."""
    required_columns = [
        'Reference designators', 
        'Quantity', 
        'Identified MPN', 
        'Identified manufacturer'
    ]

    def parse(self, file_path):
        """
        Parses the BOM file and returns a list of dictionaries, 
        each representing a BOM entry.
        """
        raise NotImplementedError("Subclasses must implement the parse method.")

    def _extract_bom_data(self, headers, rows_data):
        """Helper to extract data based on required columns and headers."""
        column_map = {}
        for col in self.required_columns:
            try:
                column_map[col] = headers.index(col)
            except ValueError:
                raise ValueError(f"Missing required column in file: {col}")

        parsed_entries = []
        for row_data in rows_data:
            # Ensure row_data has enough elements before accessing
            if len(row_data) <= max(column_map.values()):
                # print(f"Skipping malformed row (too short): {row_data}")
                continue

            try:
                mpn = str(row_data[column_map['Identified MPN']]).strip()
                manufacturer = str(row_data[column_map['Identified manufacturer']]).strip()
                
                # Try converting quantity, handle potential errors
                try:
                    quantity = int(row_data[column_map['Quantity']])
                except (ValueError, TypeError):
                    quantity = 0 # Default to 0 or raise specific error
                    # print(f"Skipping malformed row (invalid quantity): {row_data}")
                    # continue

                designators = str(row_data[column_map['Reference designators']]).strip()

                if not mpn or not manufacturer or quantity <= 0: # Skip if essential data is missing or quantity is invalid
                    continue

                parsed_entries.append({
                    'mpn': mpn,
                    'manufacturer': manufacturer,
                    'quantity': quantity,
                    'designators': designators,
                })
            except (ValueError, IndexError, TypeError) as e:
                # Log or handle rows that can't be parsed
                print(f"Skipping malformed row: {row_data} - Error: {e}")
                continue
        return parsed_entries


class XLSXParser(BaseBOMParser):
    """Parser for XLSX files."""
    def parse(self, file_path):
        try:
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active
            
            if sheet.max_row < 1:
                return [] # Empty sheet

            headers = [cell.value for cell in sheet[1]]
            rows_data = []
            for row_index in range(2, sheet.max_row + 1):
                rows_data.append([cell.value for cell in sheet[row_index]])
            
            return self._extract_bom_data(headers, rows_data)
        except Exception as e:
            raise IOError(f"Error parsing XLSX file {file_path}: {e}")


class CSVParser(BaseBOMParser):
    """Parser for CSV files."""
    def parse(self, file_path):
        try:
            with open(file_path, 'r', newline='', encoding='utf-8') as f:
                reader = csv.reader(f)
                headers = next(reader) # First row is headers
                rows_data = [list(row) for row in reader]
            
            return self._extract_bom_data(headers, rows_data)
        except Exception as e:
            raise IOError(f"Error parsing CSV file {file_path}: {e}")


class DOCXParser(BaseBOMParser):
    """Parser for DOCX files. Assumes BOM data is in the first table."""
    def parse(self, file_path):
        try:
            document = Document(file_path)
            
            if not document.tables:
                raise ValueError("No tables found in DOCX file.")
            
            # Assuming BOM data is in the first table
            table = document.tables[0]
            
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows_data = []
            for i, row in enumerate(table.rows):
                if i == 0: continue # Skip header row
                rows_data.append([cell.text.strip() for cell in row.cells])
            
            return self._extract_bom_data(headers, rows_data)
        except Exception as e:
            raise IOError(f"Error parsing DOCX file {file_path}: {e}")


class PDFParser(BaseBOMParser):
    """
    Parser for PDF files using pdfplumber. Attempts to extract tables first, then text.
    """
    def parse(self, file_path):
        try:
            with pdfplumber.open(file_path) as pdf:
                all_text_lines = []
                for page in pdf.pages:
                    # Try to extract tables first
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1: # Ensure there's a header and at least one data row
                            headers = [str(h).strip() if h is not None else '' for h in table[0]]
                            try:
                                return self._extract_bom_data(headers, table[1:])
                            except ValueError:
                                # Table headers don't match required, try next table or text
                                pass
                    
                    # If no suitable table found, extract text
                    text_content = page.extract_text()
                    if text_content:
                        all_text_lines.extend(text_content.splitlines())
                
                # Fallback to text parsing if no tables matched
                if all_text_lines:
                    # Attempt to find headers from text (similar to TXTParser)
                    lines = [line for line in all_text_lines if line.strip()]
                    if not lines: return []

                    headers = []
                    data_start_index = 0
                    for i, line in enumerate(lines):
                        # Try comma as delimiter, then space
                        potential_headers_comma = [h.strip() for h in line.split(',') if h.strip()]
                        potential_headers_space = [h.strip() for h in line.split(' ') if h.strip()] # More generic split
                        
                        if all(col in potential_headers_comma for col in self.required_columns):
                            headers = potential_headers_comma
                            data_start_index = i + 1
                            break
                        elif all(col in potential_headers_space for col in self.required_columns):
                            headers = potential_headers_space
                            data_start_index = i + 1
                            break
                    
                    if not headers:
                        raise ValueError("Could not reliably detect headers in PDF text after trying tables.")

                    rows_data = []
                    for i in range(data_start_index, len(lines)):
                        # Crude split based on detected headers/pattern
                        # This part is highly heuristic and prone to errors for PDFs without clear delimiters
                        # For now, assuming comma or space separated based on header detection attempt
                        if ',' in lines[i]:
                            rows_data.append([d.strip() for d in lines[i].split(',')])
                        else:
                            rows_data.append([d.strip() for d in lines[i].split(' ')])

                    return self._extract_bom_data(headers, rows_data)
                
            return [] # No data found
        except Exception as e:
            raise IOError(f"Error parsing PDF file {file_path}: {e}. Ensure it contains extractable text or tables.")


class TXTParser(BaseBOMParser):
    """
    Parser for TXT files. Assumes a delimited format (e.g., CSV-like with comma or tab).
    """
    def parse(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            if not lines:
                return []
            
            # Attempt to detect delimiter (tab or comma)
            first_line = lines[0]
            delimiter = ',' if ',' in first_line else ('\t' if '\t' in first_line else None)

            if not delimiter:
                raise ValueError("Could not detect a clear delimiter (comma or tab) in TXT file.")
            
            headers = [h.strip() for h in first_line.split(delimiter)]
            rows_data = []
            for line in lines[1:]:
                rows_data.append([d.strip() for d in line.split(delimiter)])
            
            return self._extract_bom_data(headers, rows_data)
        except Exception as e:
            raise IOError(f"Error parsing TXT file {file_path}: {e}")